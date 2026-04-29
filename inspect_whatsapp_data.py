import docx
import os
import argparse

def inspect_docx(path, lines=50):
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return

    doc = docx.Document(path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Print first N lines to see the structure
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] {para.text}")
            count += 1
        if count >= lines:
            break

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Quickly inspect a DOCX file structure.")
    parser.add_argument("input", help="Path to the .docx file")
    parser.add_argument("-n", "--lines", type=int, default=50, help="Number of non-empty paragraphs to show (default: 50)")
    
    args = parser.parse_args()
    inspect_docx(args.input, args.lines)
